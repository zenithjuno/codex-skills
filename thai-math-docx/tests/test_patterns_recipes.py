from __future__ import annotations

import base64
from pathlib import Path
import sys
import tempfile
import unittest

from docx import Document
from docx.oxml.ns import qn


SKILL_ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = SKILL_ROOT / "scripts"
if str(SCRIPTS) not in sys.path:
    sys.path.insert(0, str(SCRIPTS))

import thai_math_docx_patterns as patterns
import thai_math_docx_recipes as recipes


PNG_1X1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAusB9Y9ZSpMAAAAASUVORK5CYII="
)


def table_text(document) -> str:
    return "\n".join(
        paragraph.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
    )


class MaterialPatternTests(unittest.TestCase):
    def test_question_grid_is_row_major_and_uses_explicit_two_column_profile(self) -> None:
        document = Document()
        table = patterns.add_question_grid(
            document,
            [
                {"number": 1, "prompt_parts": [{"type": "text", "text": "หนึ่ง"}]},
                {"number": 2, "prompt_parts": [{"type": "text", "text": "สอง"}]},
                {"number": 3, "prompt_parts": [{"type": "text", "text": "สาม"}]},
            ],
            columns=2,
        )
        self.assertEqual("ข้อ 1. หนึ่ง", table.cell(0, 0).text)
        self.assertEqual("ข้อ 2. สอง", table.cell(0, 1).text)
        self.assertEqual("ข้อ 3. สาม", table.cell(1, 0).text)
        self.assertEqual("", table.cell(1, 1).text)
        grid = table._tbl.find(qn("w:tblGrid"))
        self.assertEqual(
            [str(round(8.5 / 2.54 * 1440))] * 2,
            [column.get(qn("w:w")) for column in grid.findall(qn("w:gridCol"))],
        )

    def test_question_grid_fails_visibly_for_unsupported_columns(self) -> None:
        with self.assertRaises(patterns.UnsupportedCapabilityError) as caught:
            patterns.add_question_grid(
                Document(),
                [{"number": 1, "prompt_parts": []}],
                columns=3,
            )
        self.assertEqual("question-grid-columns", caught.exception.capability)
        self.assertEqual(3, caught.exception.candidate["requested_columns"])

    def test_worked_example_keeps_math_editable(self) -> None:
        document = Document()
        patterns.add_worked_example(
            document,
            title="ตัวอย่าง",
            prompt_parts=[{"type": "text", "text": "จงหา "}, {"type": "math", "expr": "x"}],
            steps=[[{"type": "math", "expr": ["x", "=", "1"]}]],
        )
        self.assertIn("ตัวอย่าง", table_text(document))
        self.assertIn("<m:oMath", document.tables[0]._tbl.xml)

    def test_response_area_uses_shared_dotted_line_pattern(self) -> None:
        document = Document()
        paragraphs = patterns.add_response_area(document, label="วิธีทำ", line_count=2, dots=10)
        self.assertEqual(2, len(paragraphs))
        self.assertEqual(["." * 10, "." * 10], [paragraph.text for paragraph in paragraphs])

    def test_png_answer_visual_requires_editable_source_and_inserts(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            png = root / "answer.png"
            source = root / "answer.svg"
            png.write_bytes(PNG_1X1)
            source.write_text("<svg xmlns='http://www.w3.org/2000/svg'/>", encoding="utf-8")
            result = patterns.add_media_block(
                Document(),
                patterns.MediaBlock(png, "png-answer-visual", 3.0, editable_source_path=source),
            )
        self.assertEqual("png-answer-visual", result["media_kind"])
        self.assertTrue(result["needs_qa_review"])

    def test_svg_never_rasterizes_silently_and_expert_hook_is_marked(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            source = Path(directory) / "diagram.svg"
            source.write_text("<svg xmlns='http://www.w3.org/2000/svg'/>", encoding="utf-8")
            block = patterns.MediaBlock(source, "svg-editable", 8.0)
            with self.assertRaises(patterns.UnsupportedCapabilityError) as caught:
                patterns.add_media_block(Document(), block)
            self.assertEqual("svg-word-insertion", caught.exception.capability)

            extension = patterns.ReviewedExpertExtension(
                name="svg-package-spike",
                review_reference="REV-SVG-001",
                candidate_id="CAND-SVG-001",
                handler=lambda target, request: "inserted-by-reviewed-hook",
            )
            result = patterns.add_media_block(Document(), block, expert_extension=extension)
        self.assertEqual("svg-package-spike", result["expert_extension"])
        self.assertTrue(result["needs_qa_review"])

    def test_equations_inside_media_fail_with_candidate(self) -> None:
        with self.assertRaises(patterns.UnsupportedCapabilityError) as caught:
            patterns.add_media_block(
                Document(),
                patterns.MediaBlock(
                    "answer.png",
                    "png-answer-visual",
                    3.0,
                    editable_source_path="answer.svg",
                    contains_equations=True,
                ),
            )
        self.assertEqual("equation-inside-media", caught.exception.capability)


class FamilyRecipeTests(unittest.TestCase):
    def test_handout_recipe_assembles_shared_patterns(self) -> None:
        document = recipes.build_handout(
            title="พหุนาม",
            introduction_parts=[{"type": "text", "text": "ทบทวนแนวคิด"}],
            worked_examples=[
                {
                    "title": "ตัวอย่าง",
                    "prompt_parts": [{"type": "math", "expr": "x"}],
                    "steps": [[{"type": "math", "expr": ["x", "=", "1"]}]],
                }
            ],
            practice_questions=[
                {"number": 1, "prompt_parts": [{"type": "text", "text": "จงแยกตัวประกอบ"}]}
            ],
        )
        self.assertIn("พหุนาม", "\n".join(paragraph.text for paragraph in document.paragraphs))
        self.assertIn("จงแยกตัวประกอบ", table_text(document))

    def test_exam_recipe_keeps_written_response_area_separate(self) -> None:
        document = recipes.build_exam_paper(
            title="แนวข้อสอบ",
            instruction_parts=[{"type": "text", "text": "เลือกคำตอบ"}],
            objective_questions=[{"number": 1, "prompt_parts": [{"type": "math", "expr": "p"}]}],
            written_questions=[
                {"number": 2, "prompt_parts": [{"type": "text", "text": "แสดงวิธีทำ"}], "response_lines": 2}
            ],
        )
        paragraph_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        self.assertIn("ข้อ 2. แสดงวิธีทำ", paragraph_text)
        self.assertEqual(2, sum(paragraph.text.startswith(".") for paragraph in document.paragraphs))

    def test_answer_key_recipe_uses_editable_math(self) -> None:
        document = recipes.build_answer_key(
            title="เฉลย",
            answers=[
                {
                    "number": 1,
                    "answer_parts": [{"type": "math", "expr": ["x", "=", "1"]}],
                    "explanation_steps": [[{"type": "text", "text": "แทนค่า"}]],
                }
            ],
        )
        self.assertIn("<m:oMath", document.element.body.xml)
        self.assertIn("แทนค่า", "\n".join(paragraph.text for paragraph in document.paragraphs))


if __name__ == "__main__":
    unittest.main()

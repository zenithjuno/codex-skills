from __future__ import annotations

from pathlib import Path
import sys
import unittest

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn


SKILL_ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = SKILL_ROOT / "scripts"
if str(SCRIPTS) not in sys.path:
    sys.path.insert(0, str(SCRIPTS))

import thai_math_docx_builder as builder
import thai_math_docx_layout as layout


def width_twips(width_cm: float) -> str:
    return str(int(round(width_cm / 2.54 * 1440)))


class LayoutProfileTests(unittest.TestCase):
    def test_current_student_profile_locks_16_and_explicit_8_5_by_2(self) -> None:
        profile = layout.get_current_layout_profile("student-question-layout")
        self.assertEqual([16.0], profile["table_widths_cm"]["one-column"])
        self.assertEqual(
            [8.5, 8.5],
            profile["table_widths_cm"]["explicit-equal-two-column"],
        )
        self.assertEqual(2.54, profile["page"]["margins_cm"]["left"])

    def test_exactly_one_current_profile_per_use_case(self) -> None:
        payload = layout.load_layout_profiles()
        use_cases = {profile["use_case"] for profile in payload["profiles"]}
        for use_case in use_cases:
            current = [
                profile
                for profile in payload["profiles"]
                if profile["use_case"] == use_case and profile["status"] == "current"
            ]
            self.assertEqual(1, len(current), use_case)

    def test_logic_profile_locks_native_column_sequence(self) -> None:
        profile = layout.get_current_layout_profile("logic-practice-exam-layout")
        self.assertEqual(1, profile["sections"]["opening"]["columns"]["count"])
        self.assertEqual(2, profile["sections"]["objective"]["columns"]["count"])
        self.assertEqual(566, profile["sections"]["objective"]["columns"]["gap_twips"])
        self.assertEqual(1, profile["sections"]["written"]["columns"]["count"])
        self.assertTrue(profile["sections"]["written"]["page_break_before_content"])


class LayoutPrimitiveTests(unittest.TestCase):
    def test_cell_primitives_are_idempotent_and_keep_nil_distinct(self) -> None:
        cell = Document().add_table(rows=1, cols=1).cell(0, 0)
        for _ in range(2):
            layout.set_cell_margins(cell, top=90, start=140, bottom=100, end=150)
            layout.set_cell_shading(cell, "EAF4F8")
            layout.set_cell_borders(
                cell,
                top=layout.BorderSpec(color="59717F", size=12),
                left=None,
            )
        tc_pr = cell._tc.tcPr
        self.assertEqual(1, len(tc_pr.findall(qn("w:tcMar"))))
        self.assertEqual(1, len(tc_pr.findall(qn("w:shd"))))
        self.assertEqual("90", tc_pr.find(qn("w:tcMar")).find(qn("w:top")).get(qn("w:w")))
        borders = tc_pr.find(qn("w:tcBorders"))
        self.assertEqual("single", borders.find(qn("w:top")).get(qn("w:val")))
        self.assertEqual("nil", borders.find(qn("w:left")).get(qn("w:val")))

    def test_clear_borders_emits_all_six_explicit_nil_edges(self) -> None:
        cell = Document().add_table(rows=1, cols=1).cell(0, 0)
        layout.clear_cell_borders(cell)
        borders = cell._tc.tcPr.find(qn("w:tcBorders"))
        self.assertEqual(
            {edge: "nil" for edge in layout.CELL_BORDER_EDGES},
            {
                edge: borders.find(qn(f"w:{edge}")).get(qn("w:val"))
                for edge in layout.CELL_BORDER_EDGES
            },
        )

    def test_fixed_one_column_profile_emits_grid_total_and_cell_width(self) -> None:
        table = Document().add_table(rows=2, cols=1)
        widths = layout.apply_student_table_width_profile(table, "one-column")
        self.assertEqual([16.0], widths)
        self.assert_table_widths(table, [16.0])

    def test_explicit_equal_two_column_profile_is_not_shrunk(self) -> None:
        table = Document().add_table(rows=2, cols=2)
        widths = layout.apply_student_table_width_profile(
            table, "explicit-equal-two-column"
        )
        self.assertEqual([8.5, 8.5], widths)
        self.assert_table_widths(table, [8.5, 8.5])
        self.assertEqual(width_twips(17.0), table._tbl.tblPr.find(qn("w:tblW")).get(qn("w:w")))

    def test_merged_cells_receive_sum_of_spanned_grid_widths(self) -> None:
        table = Document().add_table(rows=2, cols=2)
        table.cell(0, 0).merge(table.cell(0, 1))
        layout.set_table_fixed_widths_cm(table, [8.5, 8.5])
        merged_tc = table.rows[0]._tr.tc_lst[0]
        self.assertEqual(width_twips(17.0), merged_tc.tcPr.find(qn("w:tcW")).get(qn("w:w")))

    def test_missing_grid_cell_fails_visibly(self) -> None:
        table = Document().add_table(rows=1, cols=2)
        table.rows[0]._tr.remove(table.rows[0]._tr.tc_lst[-1])
        with self.assertRaisesRegex(ValueError, "fewer grid columns"):
            layout.set_table_fixed_widths_cm(table, [8.5, 8.5])

    def test_repeat_header_updates_one_node(self) -> None:
        row = Document().add_table(rows=1, cols=1).rows[0]
        layout.set_repeat_table_header(row)
        layout.set_repeat_table_header(row)
        headers = row._tr.trPr.findall(qn("w:tblHeader"))
        self.assertEqual(1, len(headers))
        self.assertEqual("true", headers[0].get(qn("w:val")))

    def test_section_profile_emits_native_columns_without_separator(self) -> None:
        document = Document()
        profile = layout.get_current_layout_profile("logic-practice-exam-layout")
        section = layout.add_section_transition(
            document,
            profile["sections"]["objective"],
            start=WD_SECTION_START.CONTINUOUS,
        )
        cols = section._sectPr.find(qn("w:cols"))
        self.assertEqual("2", cols.get(qn("w:num")))
        self.assertEqual("566", cols.get(qn("w:space")))
        self.assertIsNone(cols.get(qn("w:sep")))
        self.assertEqual("continuous", section._sectPr.find(qn("w:type")).get(qn("w:val")))

    def test_return_to_one_column_removes_explicit_num(self) -> None:
        section = Document().sections[0]
        layout.set_section_columns(section, 2, gap_twips=566)
        layout.set_section_columns(section, 1, gap_twips=708)
        cols = section._sectPr.find(qn("w:cols"))
        self.assertIsNone(cols.get(qn("w:num")))
        self.assertEqual("708", cols.get(qn("w:space")))

    def test_dotted_response_line_uses_literal_periods_and_thai_16pt(self) -> None:
        document = Document()
        paragraph = layout.add_dotted_response_lines(document, count=1, dots=12)[0]
        run = paragraph.runs[0]
        self.assertEqual("." * 12, run.text)
        r_pr = run._r.rPr
        fonts = r_pr.find(qn("w:rFonts"))
        self.assertEqual("TH Sarabun New", fonts.get(qn("w:ascii")))
        self.assertEqual("TH Sarabun New", fonts.get(qn("w:cs")))
        self.assertEqual("32", r_pr.find(qn("w:szCs")).get(qn("w:val")))

    def test_layout_and_editable_omml_coexist_in_one_cell(self) -> None:
        document = builder.new_document()
        table = document.add_table(rows=1, cols=1)
        layout.apply_student_table_width_profile(table, "one-column")
        builder.append_math(table.cell(0, 0).paragraphs[0], ["x", "=", "1"])
        xml = table._tbl.xml
        self.assertIn("<m:oMath", xml)
        self.assertIn("<w:tblGrid>", xml)
        self.assertIn(f'w:w="{width_twips(16.0)}"', xml)

    def test_builder_requires_explicit_multi_column_layout(self) -> None:
        with self.assertRaisesRegex(ValueError, "explicit layout profile"):
            builder.standard_activity_table_widths(2)
        self.assertEqual(
            [8.5 / 2.54, 8.5 / 2.54],
            builder.current_student_table_widths("explicit-equal-two-column"),
        )

    def test_builder_routes_explicit_two_column_profile_to_layout_layer(self) -> None:
        table = builder.add_table(
            builder.new_document(),
            [[[], []]],
            layout_profile="explicit-equal-two-column",
        )
        self.assert_table_widths(table, [8.5, 8.5])

    def assert_table_widths(self, table, widths_cm: list[float]) -> None:
        grid = table._tbl.find(qn("w:tblGrid"))
        self.assertEqual(
            [width_twips(width) for width in widths_cm],
            [column.get(qn("w:w")) for column in grid.findall(qn("w:gridCol"))],
        )
        for row in table.rows:
            self.assertEqual(
                [width_twips(width) for width in widths_cm],
                [tc.tcPr.find(qn("w:tcW")).get(qn("w:w")) for tc in row._tr.tc_lst],
            )


if __name__ == "__main__":
    unittest.main()

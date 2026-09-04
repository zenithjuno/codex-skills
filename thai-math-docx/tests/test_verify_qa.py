from __future__ import annotations

import base64
import json
from pathlib import Path
import subprocess
import sys
import tempfile
import unittest
import zipfile

from docx import Document
from docx.oxml import parse_xml
from docx.shared import Inches


SKILL_ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = SKILL_ROOT / "scripts"
if str(SCRIPTS) not in sys.path:
    sys.path.insert(0, str(SCRIPTS))

import thai_math_docx_builder as builder
import thai_math_docx_layout as layout
import thai_math_docx_qa as qa
from thai_math_expr import frac, paren


CLI = SCRIPTS / "verify_thai_math_docx.py"
OMML_AUDIT = SCRIPTS / "audit_docx_omml.py"
PNG_1X1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAusB9Y9ZSpMAAAAASUVORK5CYII="
)


def contract(
    *,
    layout_mode: str = "standard-a4",
    media: str | dict = "none",
    source_mode: str = "generated",
    math_required: bool = False,
) -> dict:
    return {
        "schema_version": "1.0.0",
        "layout": layout_mode,
        "media": media,
        "source_mode": source_mode,
        "math": {"required": math_required},
    }


def write_contract(path: Path, payload: dict) -> Path:
    path.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")
    return path


def build_valid(path: Path, *, with_math: bool = False) -> Path:
    document = builder.new_document()
    parts = [{"type": "text", "text": "แบบฝึกหัดคณิตศาสตร์"}]
    if with_math:
        parts.extend([{"type": "text", "text": " ให้ "}, {"type": "math", "expr": ["x", "=", "1"]}])
    builder.add_paragraph(document, parts)
    return builder.save_docx(document, path)


def replace_zip_member(source: Path, output: Path, member: str, replacement: bytes) -> Path:
    with zipfile.ZipFile(source) as archive_in, zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as archive_out:
        for info in archive_in.infolist():
            data = replacement if info.filename == member else archive_in.read(info.filename)
            archive_out.writestr(info, data)
    return output


def remove_zip_member(source: Path, output: Path, member: str) -> Path:
    with zipfile.ZipFile(source) as archive_in, zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as archive_out:
        for info in archive_in.infolist():
            if info.filename != member:
                archive_out.writestr(info, archive_in.read(info.filename))
    return output


def build_math_with_png(path: Path, png: Path) -> Path:
    document = builder.new_document()
    builder.add_paragraph(document, [{"type": "math", "expr": ["x", "=", "1"]}])
    document.add_picture(str(png))
    return builder.save_docx(document, path)


class UnifiedQaCoreTests(unittest.TestCase):
    def test_malformed_contract_shape_fails_actionably(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = write_contract(
                Path(directory) / "bad.json",
                {
                    "schema_version": "1.0.0",
                    "layout": 42,
                    "media": "none",
                    "source_mode": "generated",
                },
            )
            with self.assertRaisesRegex(qa.ContractError, "layout must be a string or object"):
                qa.load_contract(path)

    def test_committed_contract_fixtures_normalize_composable_axes(self) -> None:
        payloads = {
            "standard-generated.json": contract(),
            "fixed-table-generated.json": contract(layout_mode="fixed-table", math_required=True),
            "png-answer-visual.json": contract(
                media={
                    "mode": "png-golden",
                    "role": "answer-visual",
                    "expected_count": {"min": 1, "max": 4},
                    "editability": "editable-source-required",
                    "embedding_policy": "embedded",
                    "editable_source_paths": ["editable-answer.svg"],
                }
            ),
            "custom-teacher-master.json": contract(
                layout_mode="custom-template",
                media="mixed",
                source_mode="teacher-master",
            ),
        }
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            normalized = {
                name: qa.load_contract(write_contract(root / name, payload))
                for name, payload in payloads.items()
            }
        self.assertEqual(4, len(normalized))
        self.assertEqual("standard-a4", normalized["standard-generated.json"]["layout"]["mode"])
        self.assertEqual("fixed-table", normalized["fixed-table-generated.json"]["layout"]["mode"])
        self.assertEqual("answer-visual", normalized["png-answer-visual.json"]["media"]["role"])
        self.assertEqual("teacher-master", normalized["custom-teacher-master.json"]["source_mode"])

    def test_valid_generated_docx_passes_without_final_product_claim(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = build_valid(Path(directory) / "valid.docx", with_math=True)
            result = qa.audit_docx(path, qa.load_contract(write_contract(Path(directory) / "contract.json", contract(math_required=True))))
        self.assertEqual("PASS", result["verdict"])
        self.assertFalse(result["needs_word_review"])
        self.assertIn("handoff-readiness", result["summary"])
        self.assertNotIn("final product", result["summary"].lower())
        self.assertEqual(1, result["metrics"]["omml"]["oMath_count"])

    def test_check_is_byte_for_byte_read_only(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = build_valid(Path(directory) / "valid.docx")
            before = path.read_bytes()
            result = qa.audit_docx(path, qa.load_contract(write_contract(Path(directory) / "contract.json", contract())))
            after = path.read_bytes()
        self.assertEqual("PASS", result["verdict"])
        self.assertEqual(before, after)
        self.assertTrue(result["mutation"]["source_unchanged"])
        self.assertTrue(result["mutation"]["artifact_unchanged_during_audit"])

    def test_corrupt_zip_is_artifact_fail(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "corrupt.docx"
            path.write_bytes(b"not a zip")
            result = qa.audit_docx(path, qa.load_contract(write_contract(Path(directory) / "contract.json", contract())))
        self.assertEqual("FAIL", result["verdict"])
        self.assertIn("not a valid ZIP", "\n".join(result["failures"]))

    def test_corrupt_xml_is_artifact_fail(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            source = build_valid(root / "source.docx")
            corrupt = replace_zip_member(source, root / "corrupt-xml.docx", "word/document.xml", b"<w:document>")
            result = qa.audit_docx(corrupt, qa.load_contract(write_contract(root / "contract.json", contract())))
        self.assertEqual("FAIL", result["verdict"])
        self.assertTrue(any("invalid XML" in failure for failure in result["failures"]))

    def test_font_defaults_theme_and_insertion_failures_are_combined(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "bad-fonts.docx"
            document = builder.new_document()
            paragraph = document.add_paragraph()
            unsafe = paragraph.add_run("ข้อความภาษาไทย")
            builder.set_run_font(unsafe, ascii_font="Cambria", cs_font="TH Sarabun New", size=16)
            document.save(path)
            result = qa.audit_docx(path, qa.load_contract(write_contract(Path(directory) / "contract.json", contract())))
        self.assertEqual("FAIL", result["verdict"])
        combined = "\n".join(result["failures"])
        self.assertIn("Latin w:sz=32", combined)
        self.assertIn("Thai theme font mapping", combined)

    def test_thai_inside_generic_math_fails(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "thai-math.docx"
            document = builder.new_document()
            paragraph = document.add_paragraph()
            paragraph._p.append(
                parse_xml(
                    '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" '
                    'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                    '<m:r><m:t>ภาษาไทย</m:t></m:r></m:oMath>'
                )
            )
            builder.save_docx(document, path)
            result = qa.audit_docx(path, qa.load_contract(write_contract(Path(directory) / "contract.json", contract(math_required=True))))
        self.assertEqual("FAIL", result["verdict"])
        self.assertTrue(any("generic/unformatted math" in failure for failure in result["failures"]))

    def test_full_omml_structural_audit_is_part_of_unified_qa(self) -> None:
        cases = {
            "literal-root": (["√", "18"], "literal structural glyph '√'"),
            "redundant-radicand-delimiter": (
                {"kind": "rad", "deg": ["3"], "items": [paren(["−", "64"])]},
                "entire radicand is wrapped",
            ),
            "numerator-unary-minus": (
                frac(["−", "1"], "4"),
                "unary minus starts inside fraction numerator",
            ),
            "literal-set-braces": (
                ["{", "−", "3", ",", "4", "}"],
                "opening set brace",
            ),
        }
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            qa_contract = qa.load_contract(
                write_contract(root / "contract.json", contract(math_required=True))
            )
            for name, (value, expected) in cases.items():
                with self.subTest(name=name):
                    document = builder.new_document()
                    builder.add_paragraph(document, [{"type": "math", "expr": value}])
                    path = builder.save_docx(document, root / f"{name}.docx")
                    result = qa.audit_docx(path, qa_contract)
                    self.assertEqual("FAIL", result["verdict"], result)
                    self.assertIn(expected, "\n".join(result["failures"]))

            valid = builder.new_document()
            builder.add_paragraph(
                valid,
                [
                    {
                        "type": "math",
                        "expr": [
                            "−",
                            frac("1", "4"),
                            "+",
                            {"kind": "rad", "items": ["18"]},
                            "+",
                            paren(["−", "3", ",", "4"], beg="{", end="}"),
                        ],
                    }
                ],
            )
            valid_path = builder.save_docx(valid, root / "valid-structures.docx")
            valid_result = qa.audit_docx(valid_path, qa_contract)
        self.assertEqual("PASS", valid_result["verdict"], valid_result)

    def test_upright_fused_coefficient_variable_fails_unified_qa(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            document = builder.new_document()
            paragraph = document.add_paragraph()
            paragraph._p.append(
                parse_xml(
                    '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
                    '<m:r><m:rPr><m:nor/></m:rPr><m:t>3x</m:t></m:r>'
                    '</m:oMath>'
                )
            )
            builder.append_parts(paragraph, [{"type": "text", "text": " "}])
            path = builder.save_docx(document, root / "fused.docx")
            result = qa.audit_docx(
                path,
                qa.load_contract(
                    write_contract(root / "contract.json", contract(math_required=True))
                ),
            )
        self.assertEqual("FAIL", result["verdict"])
        self.assertIn("fuses a coefficient to a variable", "\n".join(result["failures"]))

    def test_fixed_table_contract_checks_grid_and_missing_cells(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            good = builder.new_document()
            builder.add_table(good, [[[{"type": "text", "text": "ข้อ"}]]], layout_profile="one-column")
            good_path = builder.save_docx(good, root / "fixed-good.docx")
            fixed_contract = qa.load_contract(write_contract(root / "contract.json", contract(layout_mode="fixed-table")))
            self.assertEqual("PASS", qa.audit_docx(good_path, fixed_contract)["verdict"])

            bad = builder.new_document()
            table = bad.add_table(rows=1, cols=2)
            layout.set_table_fixed_widths_cm(table, [8.5, 8.5])
            table.rows[0]._tr.remove(table.rows[0]._tr.tc_lst[-1])
            bad_path = builder.save_docx(bad, root / "fixed-bad.docx")
            bad_result = qa.audit_docx(bad_path, fixed_contract)
        self.assertEqual("FAIL", bad_result["verdict"])
        self.assertTrue(any("row occupies" in failure for failure in bad_result["failures"]))

    def test_native_columns_contract_checks_real_word_columns(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            document = builder.new_document()
            profile = layout.get_current_layout_profile("logic-practice-exam-layout")
            layout.add_section_transition(document, profile["sections"]["objective"])
            path = builder.save_docx(document, root / "columns.docx")
            result = qa.audit_docx(path, qa.load_contract(write_contract(root / "contract.json", contract(layout_mode="native-columns"))))
        self.assertEqual("PASS", result["verdict"])
        self.assertTrue(any(section["columns"] == 2 for section in result["metrics"]["layout"]["sections"]))

    def test_unknown_media_is_review_only_under_mixed_contract(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            path = build_valid(root / "media.docx")
            with zipfile.ZipFile(path, "a", zipfile.ZIP_DEFLATED) as archive:
                archive.writestr("word/media/unknown.bin", b"opaque")
            result = qa.audit_docx(path, qa.load_contract(write_contract(root / "contract.json", contract(media="mixed"))))
        self.assertEqual("PASS", result["verdict"])
        self.assertTrue(result["needs_word_review"])
        self.assertTrue(any("unknown media" in item for item in result["review_items"]))

    def test_media_contract_violation_fails(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            path = build_valid(root / "media.docx")
            with zipfile.ZipFile(path, "a", zipfile.ZIP_DEFLATED) as archive:
                archive.writestr("word/media/unknown.bin", b"opaque")
            result = qa.audit_docx(path, qa.load_contract(write_contract(root / "contract.json", contract(media="none"))))
        self.assertEqual("FAIL", result["verdict"])
        self.assertTrue(any("contract is none" in failure for failure in result["failures"]))

    def test_png_media_contract_carries_role_count_editability_and_embedding(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            png = root / "answer.png"
            editable = root / "answer.svg"
            png.write_bytes(PNG_1X1)
            editable.write_text("<svg xmlns='http://www.w3.org/2000/svg'/>", encoding="utf-8")
            document = builder.new_document()
            document.add_picture(str(png))
            path = builder.save_docx(document, root / "png.docx")
            media_contract = {
                "mode": "png-golden",
                "role": "answer-visual",
                "expected_count": {"min": 1, "max": 1},
                "editability": "editable-source-required",
                "embedding_policy": "embedded",
                "editable_source_paths": [str(editable)],
            }
            result = qa.audit_docx(
                path,
                qa.load_contract(write_contract(root / "contract.json", contract(media=media_contract))),
            )
        self.assertEqual("PASS", result["verdict"])
        self.assertTrue(result["needs_word_review"])
        self.assertEqual("answer-visual", result["metrics"]["media"]["role"])
        self.assertEqual(1, result["metrics"]["media"]["count"])
        self.assertEqual("editable-source-required", result["metrics"]["media"]["editability"])
        self.assertEqual("embedded", result["metrics"]["media"]["embedding_policy"])

    def test_allowed_image_does_not_fail_omml_audit(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            png = root / "diagram.png"
            png.write_bytes(PNG_1X1)
            path = build_math_with_png(root / "math-with-image.docx", png)
            completed = subprocess.run(
                [sys.executable, str(OMML_AUDIT), str(path)],
                capture_output=True,
                text=True,
                check=False,
            )
        self.assertEqual(0, completed.returncode, completed.stdout + completed.stderr)
        self.assertIn("image_count: 1", completed.stdout)
        self.assertIn("image validity is handled by the media QA contract", completed.stdout)

    def test_valid_internal_image_relationship_passes_media_qa(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            png = root / "diagram.png"
            png.write_bytes(PNG_1X1)
            path = build_math_with_png(root / "math-with-image.docx", png)
            result = qa.audit_docx(
                path,
                qa.load_contract(write_contract(root / "contract.json", contract(media="mixed", math_required=True))),
            )
        self.assertEqual("PASS", result["verdict"])
        self.assertEqual(1, result["metrics"]["media"]["internal_image_relationship_count"])
        self.assertEqual([], result["metrics"]["media"]["broken_internal_image_relationships"])

    def test_broken_internal_image_relationship_fails_media_qa(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            png = root / "diagram.png"
            png.write_bytes(PNG_1X1)
            source = build_math_with_png(root / "source.docx", png)
            with zipfile.ZipFile(source) as archive:
                image_part = next(name for name in archive.namelist() if name.startswith("word/media/"))
            broken = remove_zip_member(source, root / "broken.docx", image_part)
            result = qa.audit_docx(
                broken,
                qa.load_contract(write_contract(root / "contract.json", contract(media="mixed", math_required=True))),
            )
        self.assertEqual("FAIL", result["verdict"])
        self.assertTrue(any("broken internal image relationship" in item for item in result["failures"]))

    def test_custom_or_imported_source_sets_independent_review_flag(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            path = build_valid(root / "imported.docx")
            result = qa.audit_docx(
                path,
                qa.load_contract(write_contract(root / "contract.json", contract(layout_mode="custom-template", source_mode="teacher-master"))),
            )
        self.assertEqual("PASS", result["verdict"])
        self.assertTrue(result["needs_word_review"])
        self.assertGreaterEqual(len(result["review_items"]), 2)

    def test_fix_and_check_writes_new_output_and_preserves_source(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            source = root / "source.docx"
            document = Document()
            section = document.sections[0]
            section.page_width = Inches(8.27)
            section.page_height = Inches(11.69)
            section.top_margin = Inches(1)
            section.right_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            document.add_paragraph("ข้อความ")
            document.save(source)
            before = source.read_bytes()
            output, source_hash = qa.fix_copy(source, root / "fixed.docx")
            result = qa.audit_docx(
                output,
                qa.load_contract(write_contract(root / "contract.json", contract())),
                mode="fix-and-check",
                source_path=source,
                source_sha256_before=source_hash,
            )
            self.assertEqual(before, source.read_bytes())
            self.assertEqual("PASS", result["verdict"])
            self.assertTrue(result["mutation"]["source_unchanged"])
            self.assertNotEqual(result["mutation"]["artifact_sha256_after"], source_hash)

    def test_fix_refuses_source_overwrite(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = build_valid(Path(directory) / "source.docx")
            with self.assertRaisesRegex(qa.ContractError, "must not overwrite"):
                qa.fix_copy(path, path)

    def test_reports_always_write_json_and_explicit_dir_adds_markdown(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            path = build_valid(root / "valid.docx")
            result = qa.audit_docx(path, qa.load_contract(write_contract(root / "contract.json", contract())))
            reports = qa.write_reports(result, report_dir=root / "reports")
            loaded = json.loads(Path(reports["json"]).read_text(encoding="utf-8"))
        self.assertEqual("PASS", loaded["verdict"])
        self.assertTrue(Path(reports["markdown"]).name.endswith(".qa.md"))
        self.assertIn("handoff-readiness", loaded["summary"])


class UnifiedQaCliTests(unittest.TestCase):
    def run_cli(self, cwd: Path, *args: str) -> subprocess.CompletedProcess[str]:
        return subprocess.run(
            [sys.executable, str(CLI), *args],
            cwd=cwd,
            capture_output=True,
            text=True,
            check=False,
        )

    def test_exact_exit_codes_and_short_examples(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            valid = build_valid(root / "valid.docx")
            corrupt = root / "corrupt.docx"
            corrupt.write_bytes(b"bad")
            contract_path = write_contract(root / "contract.json", contract())
            passed = self.run_cli(root, "check", str(valid), "--contract", str(contract_path))
            failed = self.run_cli(root, "check", str(corrupt), "--contract", str(contract_path))
            blocked = self.run_cli(root, "check", str(root / "missing.docx"), "--contract", str(contract_path))
            report = root / "qa-reports/valid.qa.json"
            loaded = json.loads(report.read_text(encoding="utf-8"))
        self.assertEqual(0, passed.returncode, passed.stdout + passed.stderr)
        self.assertIn("PASS: automated handoff-readiness", passed.stdout)
        self.assertEqual(1, failed.returncode, failed.stdout + failed.stderr)
        self.assertIn("FAIL: automated handoff-readiness", failed.stdout)
        self.assertEqual(2, blocked.returncode, blocked.stdout + blocked.stderr)
        self.assertIn("BLOCKED", blocked.stdout)
        self.assertEqual("PASS", loaded["verdict"])

    def test_invalid_contract_is_blocked(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            path = build_valid(root / "valid.docx")
            invalid = write_contract(root / "invalid.json", {"schema_version": "9"})
            result = self.run_cli(root, "check", str(path), "--contract", str(invalid))
        self.assertEqual(2, result.returncode)
        self.assertIn("schema_version", result.stdout)


class SingleCommandGateTests(unittest.TestCase):
    """One command must cover the whole document.

    Anything a worker could forget to run belongs inside the gate. The one
    deliberate exception (CHG-001, thai-docx seam) is the plain-text-math scan:
    it false-fails on ordinary prose relations (e.g. "คะแนน ≥ 80"), so it runs
    only when the document is a math document (math.required, or equations
    present) and is legitimately absent for a declared math-free document.
    """

    ALWAYS = ("package-xml-integrity", "thai-fonts-defaults-theme-insertion",
              "omml-editability", "geometry-table-shape", "media-contract",
              "mutation-provenance")

    def _audit(self, math_required: bool):
        import tempfile
        import thai_math_docx_builder as builder
        import thai_math_docx_qa as qa
        with tempfile.TemporaryDirectory() as tmp:
            doc = builder.new_document()
            builder.add_paragraph(doc, [{"type": "text", "text": "ทดสอบ"}])
            path = builder.save_docx(doc, Path(tmp) / "gate.docx")
            contract = qa.normalize_contract({
                "schema_version": "1.0.0", "layout": "standard-a4", "media": "none",
                "source_mode": "generated", "math": {"required": math_required},
            })
            result = qa.audit_docx(path, contract, mode="check")
        return {check["id"] for check in result["checks"]}

    def test_gate_reports_every_document_check_for_a_math_document(self) -> None:
        ids = self._audit(math_required=True)
        for expected in self.ALWAYS:
            self.assertIn(expected, ids)
        self.assertIn("math-in-plain-text", ids)

    def test_math_free_document_runs_every_check_except_the_gated_plain_text_scan(self) -> None:
        ids = self._audit(math_required=False)
        for expected in self.ALWAYS:
            self.assertIn(expected, ids)
        # CHG-001: the plain-text-math scan is gated off for a declared math-free doc.
        self.assertNotIn("math-in-plain-text", ids)

    def test_relational_maths_in_plain_text_fails_the_gate(self) -> None:
        import tempfile
        import thai_math_docx_builder as builder
        import thai_math_docx_qa as qa
        with tempfile.TemporaryDirectory() as tmp:
            doc = builder.new_document()
            builder.add_paragraph(doc, [{"type": "text", "text": "ถ้า x = 0 แล้วจบ"}])
            path = builder.save_docx(doc, Path(tmp) / "leak.docx")
            result = qa.audit_docx(path, qa.load_contract(None), mode="check")
        self.assertEqual("FAIL", result["verdict"])
        self.assertTrue(any("fused into one run" in f for f in result["failures"]), result["failures"])


if __name__ == "__main__":
    unittest.main()

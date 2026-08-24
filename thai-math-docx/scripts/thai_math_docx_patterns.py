#!/usr/bin/env python3
"""Reusable material patterns above the Thai math builder and layout layers."""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import re
from typing import Any, Callable, Mapping, Sequence
import xml.etree.ElementTree as ET
from xml.sax.saxutils import escape

from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.part import Part
from docx.shared import Cm

import thai_math_docx_builder as builder
import thai_math_docx_layout as layout


class UnsupportedCapabilityError(RuntimeError):
    """Visible failure carrying a candidate record for the work-batch review."""

    def __init__(self, capability: str, reason: str, candidate: Mapping[str, Any]):
        super().__init__(
            f"Unsupported capability {capability!r}: {reason}. "
            "Record the attached candidate and run the relevant QA review; do not approximate."
        )
        self.capability = capability
        self.candidate = dict(candidate)


@dataclass(frozen=True)
class ReviewedExpertExtension:
    name: str
    review_reference: str
    candidate_id: str
    handler: Callable[[Any, Mapping[str, Any]], Any]

    def apply(self, target: Any, request: Mapping[str, Any]) -> dict[str, Any]:
        if not self.review_reference.strip() or not self.candidate_id.strip():
            raise ValueError("expert extensions require a review_reference and candidate_id")
        return {
            "result": self.handler(target, request),
            "expert_extension": self.name,
            "review_reference": self.review_reference,
            "candidate_id": self.candidate_id,
            "needs_qa_review": True,
        }


@dataclass(frozen=True)
class MediaBlock:
    source_path: str | Path
    media_kind: str
    width_cm: float
    editable_source_path: str | Path | None = None
    contains_equations: bool = False


def _svg_aspect_ratio(source: Path) -> float:
    root = ET.fromstring(source.read_bytes())
    view_box = root.get("viewBox")
    if view_box:
        values = [float(value) for value in re.split(r"[\s,]+", view_box.strip())]
        if len(values) == 4 and values[2] > 0 and values[3] > 0:
            return values[2] / values[3]

    def numeric_dimension(name: str) -> float | None:
        value = root.get(name)
        if not value:
            return None
        match = re.fullmatch(r"\s*([0-9]+(?:\.[0-9]+)?)\s*(?:px|pt|pc|mm|cm|in)?\s*", value)
        return float(match.group(1)) if match else None

    width = numeric_dimension("width")
    height = numeric_dimension("height")
    if width and height and width > 0 and height > 0:
        return width / height
    raise ValueError(f"SVG requires a positive viewBox or numeric width and height: {source}")


def add_svg_picture(
    container: Any,
    source_path: str | Path,
    *,
    width_cm: float,
    alt_text: str | None = None,
) -> Any:
    """Embed an SVG package part as a centered inline Word picture."""
    source = Path(source_path)
    if source.suffix.lower() != ".svg" or not source.is_file():
        raise ValueError(f"SVG source must be an existing .svg file: {source}")
    if width_cm <= 0:
        raise ValueError("SVG width_cm must be greater than zero")

    paragraph = container.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    package = paragraph.part.package
    svg_part = Part(
        package.next_partname("/word/media/image%d.svg"),
        "image/svg+xml",
        source.read_bytes(),
        package,
    )
    relationship_id = paragraph.part.relate_to(svg_part, RT.IMAGE)

    doc_pr_ids = [
        int(node.get("id", "0"))
        for node in paragraph.part.element.iter(qn("wp:docPr"))
        if node.get("id", "0").isdigit()
    ]
    drawing_id = max(doc_pr_ids, default=0) + 1
    width_emu = int(Cm(width_cm))
    height_emu = round(width_emu / _svg_aspect_ratio(source))
    safe_name = escape(source.name, {'"': "&quot;"})
    safe_description = escape(alt_text or source.stem, {'"': "&quot;"})

    drawing = parse_xml(
        f"""
        <w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
          xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
          xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
          xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"
          xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
          <wp:inline distT="0" distB="0" distL="0" distR="0">
            <wp:extent cx="{width_emu}" cy="{height_emu}"/>
            <wp:effectExtent l="0" t="0" r="0" b="0"/>
            <wp:docPr id="{drawing_id}" name="{safe_name}" descr="{safe_description}"/>
            <wp:cNvGraphicFramePr><a:graphicFrameLocks noChangeAspect="1"/></wp:cNvGraphicFramePr>
            <a:graphic>
              <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
                <pic:pic>
                  <pic:nvPicPr>
                    <pic:cNvPr id="0" name="{safe_name}"/>
                    <pic:cNvPicPr/>
                  </pic:nvPicPr>
                  <pic:blipFill>
                    <a:blip r:embed="{relationship_id}"/>
                    <a:stretch><a:fillRect/></a:stretch>
                  </pic:blipFill>
                  <pic:spPr>
                    <a:xfrm>
                      <a:off x="0" y="0"/>
                      <a:ext cx="{width_emu}" cy="{height_emu}"/>
                    </a:xfrm>
                    <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
                    <a:noFill/>
                    <a:ln><a:noFill/></a:ln>
                  </pic:spPr>
                </pic:pic>
              </a:graphicData>
            </a:graphic>
          </wp:inline>
        </w:drawing>
        """
    )
    paragraph.add_run()._r.append(drawing)
    return paragraph


def _validate_parts(parts: Sequence[Mapping[str, Any]], field: str) -> None:
    for part in parts:
        if part.get("type") == "table":
            raise ValueError(f"{field} cannot contain a nested table part")


def add_question_grid(
    document: Any,
    questions: Sequence[Mapping[str, Any]],
    *,
    columns: int = 1,
    cell_margins_twips: Mapping[str, int] | None = None,
) -> Any:
    """Add a fixed, borderless row-major question grid using current profiles."""
    if columns not in (1, 2):
        raise UnsupportedCapabilityError(
            "question-grid-columns",
            f"current shared pattern supports 1 or explicit equal 2 columns, got {columns}",
            {"capability": "question-grid-columns", "requested_columns": columns},
        )
    if not questions:
        raise ValueError("questions must not be empty")
    for question in questions:
        if "number" not in question or "prompt_parts" not in question:
            raise ValueError("every question requires number and prompt_parts")
        _validate_parts(question["prompt_parts"], "prompt_parts")

    row_count = (len(questions) + columns - 1) // columns
    table = document.add_table(rows=row_count, cols=columns)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    profile = "one-column" if columns == 1 else "explicit-equal-two-column"
    layout.apply_student_table_width_profile(table, profile)
    margins = dict(cell_margins_twips or {"top": 70, "start": 100, "bottom": 70, "end": 100})

    for index in range(row_count * columns):
        cell = table.cell(index // columns, index % columns)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        layout.clear_cell_borders(cell)
        layout.set_cell_margins(cell, **margins)
        paragraph = cell.paragraphs[0]
        builder.configure_paragraph(paragraph, space_after=0)
        if index >= len(questions):
            continue
        question = questions[index]
        label = paragraph.add_run(f"ข้อ {question['number']}. ")
        builder.set_thai_label_run(label, bold=True)
        builder.append_parts(paragraph, list(question["prompt_parts"]))
        if question.get("response_lines"):
            layout.add_dotted_response_lines(
                cell,
                count=int(question["response_lines"]),
                dots=int(question.get("response_dots", 60)),
            )
    return table


def add_worked_example(
    document: Any,
    *,
    title: str,
    prompt_parts: Sequence[Mapping[str, Any]],
    steps: Sequence[Sequence[Mapping[str, Any]]],
    heading_fill: str = "EAF4F8",
) -> Any:
    """Add one current 16 cm worked-example box with semantic prompt and steps."""
    _validate_parts(prompt_parts, "prompt_parts")
    if not steps:
        raise ValueError("worked example requires at least one step")
    for step in steps:
        _validate_parts(step, "steps")
    table = document.add_table(rows=2, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    layout.apply_student_table_width_profile(table, "one-column")
    layout.set_cell_shading(table.cell(0, 0), heading_fill)
    layout.set_cell_margins(table.cell(0, 0), top=70, start=120, bottom=70, end=120)
    layout.set_cell_margins(table.cell(1, 0), top=90, start=120, bottom=90, end=120)

    heading = table.cell(0, 0).paragraphs[0]
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    builder.configure_paragraph(heading, space_after=0)
    builder.set_thai_label_run(heading.add_run(title), bold=True)

    body = table.cell(1, 0)
    prompt = body.paragraphs[0]
    builder.configure_paragraph(prompt, space_after=4)
    builder.append_parts(prompt, list(prompt_parts))
    for number, step in enumerate(steps, start=1):
        paragraph = body.add_paragraph()
        builder.configure_paragraph(paragraph, space_after=2)
        builder.set_thai_label_run(paragraph.add_run(f"ขั้นที่ {number}  "), bold=True)
        builder.append_parts(paragraph, list(step))
    return table


def add_response_area(
    container: Any,
    *,
    label: str | None = None,
    line_count: int = 3,
    dots: int = 80,
) -> list[Any]:
    if label:
        paragraph = container.add_paragraph()
        builder.configure_paragraph(paragraph, space_after=0)
        builder.set_thai_label_run(paragraph.add_run(label), bold=True)
    return layout.add_dotted_response_lines(container, count=line_count, dots=dots)


def add_media_block(
    container: Any,
    block: MediaBlock,
    *,
    expert_extension: ReviewedExpertExtension | None = None,
) -> dict[str, Any]:
    """Insert an approved media block or fail with a reviewable candidate."""
    source = Path(block.source_path)
    if block.width_cm <= 0:
        raise ValueError("media width_cm must be greater than zero")
    if block.contains_equations:
        raise UnsupportedCapabilityError(
            "equation-inside-media",
            "equations must remain editable OMML rather than image content",
            {"capability": "equation-inside-media", "source_path": str(source)},
        )
    if not source.is_file():
        raise FileNotFoundError(f"media source does not exist: {source}")
    if block.media_kind == "png-answer-visual":
        if source.suffix.lower() != ".png" or not block.editable_source_path:
            raise ValueError("PNG answer visuals require a .png plus editable_source_path")
        editable_source = Path(block.editable_source_path)
        if not editable_source.is_file():
            raise FileNotFoundError(f"editable media source does not exist: {editable_source}")
        paragraph = container.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(str(source), width=Cm(block.width_cm))
        return {
            "media_kind": block.media_kind,
            "source_path": str(source),
            "editable_source_path": str(editable_source),
            "needs_qa_review": True,
        }
    if block.media_kind == "svg-editable":
        if source.suffix.lower() != ".svg":
            raise ValueError("svg-editable media requires an .svg source")
        if expert_extension is not None:
            return expert_extension.apply(container, {"block": block})
        add_svg_picture(
            container,
            source,
            width_cm=block.width_cm,
            alt_text=source.stem,
        )
        return {
            "media_kind": block.media_kind,
            "source_path": str(source),
            "needs_qa_review": True,
        }
    raise UnsupportedCapabilityError(
        "media-kind",
        f"unknown media kind {block.media_kind!r}",
        {"capability": "media-kind", "requested_kind": block.media_kind},
    )

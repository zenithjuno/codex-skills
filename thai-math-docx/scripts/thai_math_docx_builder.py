#!/usr/bin/env python3
"""Shared Thai math DOCX builder primitives.

Use this module as the starting point for generated Thai mathematics DOCX
scripts. It is intentionally not tied to JSON. Call ``append_parts`` with
Python dict/list parts, or call the lower-level run/math helpers directly.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any
import tempfile
import zipfile
import xml.etree.ElementTree as ET
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from thai_math_docx_layout import get_current_layout_profile, set_table_fixed_widths_cm
from thai_math_source_adapter import normalize_math_string


M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"

# Teacher-confirmed standard for equal-column student activity/answer tables.
# These values are a fixed table-grid target; data tables may deliberately use
# a different, unequal width specification.
STANDARD_ACTIVITY_TABLE_WIDTH_CM = 16.0
STANDARD_ACTIVITY_TABLE_WIDTH_IN = STANDARD_ACTIVITY_TABLE_WIDTH_CM / 2.54

VARIABLES = {
    "A", "B", "C", "D", "E", "F", "G", "H", "I", "K", "L", "M", "N", "P", "Q", "R",
    "S", "T", "U", "V", "W", "X", "Y", "Z",
    "a", "b", "c", "d", "e", "f", "g", "h", "i", "j", "k", "m", "n", "p", "q", "r",
    "s", "t", "u", "v", "w", "x", "y", "z",
    "α", "β", "γ", "θ", "μ", "π", "σ",
}

FUNCTION_NAMES = {
    "sin", "cos", "tan", "sec", "cosec", "csc", "cot", "log", "ln", "lim", "det", "adj",
    "arg", "Re", "Im",
}

OPERATOR_TOKENS = {
    "+", "−", "-", "=", "<", ">", "≤", "≥", "≠", "∈", "∉", "⊂", "⊆", "∪", "∩",
    "∧", "∨", "↔", "→", "⇒", "⇔", "*", "×", "·", ":", "!", "%", "|", "∣", "∘",
}

TIGHT_PREFIX_OPERATORS = {"~", "¬", "∀", "∃"}


def ensure_child(parent: Any, tag: str) -> Any:
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def contains_thai(text: str) -> bool:
    return any("\u0e00" <= ch <= "\u0e7f" for ch in str(text))


def set_run_font(run: Any, ascii_font: str = "Cambria", cs_font: str = "TH Sarabun New", size: int = 12) -> None:
    run.font.name = ascii_font
    run.font.size = Pt(size)
    r_pr = run._r.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), ascii_font)
    r_fonts.set(qn("w:hAnsi"), ascii_font)
    r_fonts.set(qn("w:cs"), cs_font)


def set_default_run_properties(r_pr: Any) -> None:
    r_fonts = ensure_child(r_pr, "w:rFonts")
    for attr in ("w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"):
        r_fonts.attrib.pop(qn(attr), None)
    r_fonts.set(qn("w:ascii"), "Cambria")
    r_fonts.set(qn("w:hAnsi"), "Cambria")
    r_fonts.set(qn("w:cs"), "TH Sarabun New")
    ensure_child(r_pr, "w:sz").set(qn("w:val"), "24")
    ensure_child(r_pr, "w:szCs").set(qn("w:val"), "32")
    lang = ensure_child(r_pr, "w:lang")
    lang.set(qn("w:val"), "en-US")
    lang.set(qn("w:bidi"), "th-TH")


def enforce_document_font_defaults(doc: Document) -> None:
    styles = doc.styles.element
    doc_defaults = styles.find(qn("w:docDefaults"))
    if doc_defaults is None:
        doc_defaults = OxmlElement("w:docDefaults")
        styles.insert(0, doc_defaults)
    r_pr_default = ensure_child(doc_defaults, "w:rPrDefault")
    set_default_run_properties(ensure_child(r_pr_default, "w:rPr"))
    set_default_run_properties(doc.styles["Normal"]._element.get_or_add_rPr())


def set_thai_body_run(run: Any, bold: bool | None = None, size: int = 16) -> None:
    """Insertion-safe Thai body run: visible Thai 16pt, future Latin 12pt."""
    if bold is not None:
        run.bold = bold
    set_run_font(run, ascii_font="Cambria", cs_font="TH Sarabun New", size=12)
    r_pr = run._r.get_or_add_rPr()
    ensure_child(r_pr, "w:sz").set(qn("w:val"), "24")
    ensure_child(r_pr, "w:szCs").set(qn("w:val"), str(size * 2))
    if r_pr.find(qn("w:cs")) is None:
        r_pr.append(OxmlElement("w:cs"))
    lang = ensure_child(r_pr, "w:lang")
    lang.set(qn("w:val"), "en-US")
    lang.set(qn("w:bidi"), "th-TH")


def set_thai_label_run(run: Any, bold: bool | None = None, size: int = 16) -> None:
    """All-slot Thai label run for question numbers, choice labels, headings."""
    if bold is not None:
        run.bold = bold
    set_run_font(run, ascii_font="TH Sarabun New", cs_font="TH Sarabun New", size=size)
    r_pr = run._r.get_or_add_rPr()
    ensure_child(r_pr, "w:sz").set(qn("w:val"), str(size * 2))
    ensure_child(r_pr, "w:szCs").set(qn("w:val"), str(size * 2))
    if r_pr.find(qn("w:cs")) is None:
        r_pr.append(OxmlElement("w:cs"))
    lang = ensure_child(r_pr, "w:lang")
    lang.set(qn("w:val"), "th-TH")
    lang.set(qn("w:bidi"), "th-TH")


def set_latin_run(run: Any, bold: bool | None = None, size: int = 12) -> None:
    if bold is not None:
        run.bold = bold
    set_run_font(run, ascii_font="Cambria", cs_font="TH Sarabun New", size=size)
    r_pr = run._r.get_or_add_rPr()
    ensure_child(r_pr, "w:sz").set(qn("w:val"), str(size * 2))
    lang = ensure_child(r_pr, "w:lang")
    lang.set(qn("w:val"), "en-US")
    lang.attrib.pop(qn("w:bidi"), None)


def configure_paragraph(paragraph: Any, space_after: int = 4) -> None:
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(space_after)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    enforce_document_font_defaults(doc)
    normal = doc.styles["Normal"]
    normal.font.name = "Cambria"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal.paragraph_format.line_spacing = 1.0
    normal.paragraph_format.space_after = Pt(4)


def new_document() -> Document:
    doc = Document()
    configure_document(doc)
    return doc


def math_run(text: Any, m_rpr: str = "", preserve_space: bool = False) -> str:
    """Latin/symbol OMML run.

    ``w:szCs`` stays at 32 (16 pt) even though math glyphs never use the
    Complex Script slot: Word formats text typed straight after an equation
    from that equation's last run, so a smaller ``szCs`` here would silently
    drop manually typed Thai to 12 pt.
    """
    text = escape(str(text))
    space_attr = ' xml:space="preserve"' if preserve_space else ""
    return (
        "<m:r>"
        f"{m_rpr}"
        '<w:rPr><w:sz w:val="24"/><w:szCs w:val="32"/></w:rPr>'
        f"<m:t{space_attr}>{text}</m:t>"
        "</m:r>"
    )


def mr(text: Any, italic: bool = True) -> str:
    sty = '<m:rPr><m:sty m:val="i"/></m:rPr>' if italic else ""
    return math_run(text, sty)


def mtext(text: Any) -> str:
    raw = str(text)
    return math_run(raw, "<m:rPr><m:nor/></m:rPr>", raw != raw.strip())


def mop(text: Any) -> str:
    return math_run(str(text))


def thai_mtext(text: Any) -> str:
    raw = str(text)
    text_xml = escape(raw)
    space_attr = ' xml:space="preserve"' if raw != raw.strip() else ""
    return (
        "<m:r>"
        "<m:rPr><m:nor/></m:rPr>"
        "<w:rPr>"
        '<w:rFonts w:ascii="TH Sarabun New" w:hAnsi="TH Sarabun New" w:cs="TH Sarabun New"/>'
        '<w:sz w:val="32"/><w:szCs w:val="32"/>'
        "<w:cs/>"
        '<w:lang w:val="th-TH" w:bidi="th-TH"/>'
        "</w:rPr>"
        f"<m:t{space_attr}>{text_xml}</m:t>"
        "</m:r>"
    )


def items_fragment(items: list[Any]) -> str:
    return "".join(item_to_omml_fragment(item) for item in items)


def delim(items: list[Any], beg_chr: str = "(", end_chr: str = ")") -> str:
    if beg_chr == "{" and end_chr == "}":
        return mtext("{") + items_fragment(items) + mtext("}")
    return (
        f'<m:d><m:dPr><m:begChr m:val="{escape(beg_chr)}"/><m:endChr m:val="{escape(end_chr)}"/></m:dPr>'
        "<m:e>" + items_fragment(items) + "</m:e></m:d>"
    )


def matrix_column_properties(col_aligns: list[str] | None = None, base_jc: str | None = None) -> str:
    pieces = []
    if base_jc:
        pieces.append(f'<m:baseJc m:val="{escape(base_jc)}"/>')
    if col_aligns:
        cols = []
        for align in col_aligns:
            cols.append(
                "<m:mc><m:mcPr>"
                '<m:count m:val="1"/>'
                f'<m:mcJc m:val="{escape(align)}"/>'
                "</m:mcPr></m:mc>"
            )
        pieces.append("<m:mcs>" + "".join(cols) + "</m:mcs>")
    return "<m:mPr>" + "".join(pieces) + "</m:mPr>" if pieces else ""


def matrix_fragment(
    rows: list[list[list[Any]]],
    col_aligns: list[str] | None = None,
    base_jc: str | None = None,
) -> str:
    row_fragments = []
    for row in rows:
        cells = "".join("<m:e>" + items_fragment(cell) + "</m:e>" for cell in row)
        row_fragments.append("<m:mr>" + cells + "</m:mr>")
    return "<m:m>" + matrix_column_properties(col_aligns, base_jc) + "".join(row_fragments) + "</m:m>"


def bracketed_matrix_fragment(
    rows: list[list[list[Any]]],
    brackets: str | tuple[str, str] | None = "[]",
    col_aligns: list[str] | None = None,
    base_jc: str | None = None,
) -> str:
    matrix = matrix_fragment(rows, col_aligns, base_jc)
    if brackets in (None, "", "none"):
        return matrix
    bracket_pairs = {
        "[]": ("[", "]"),
        "()": ("(", ")"),
        "||": ("|", "|"),
        "{}": ("{", "}"),
    }
    if isinstance(brackets, str):
        beg, end = bracket_pairs.get(brackets, ("[", "]"))
    else:
        beg, end = brackets
    return (
        f'<m:d><m:dPr><m:begChr m:val="{escape(beg)}"/><m:endChr m:val="{escape(end)}"/></m:dPr>'
        "<m:e>" + matrix + "</m:e></m:d>"
    )


def accent_character(expr: dict[str, Any]) -> str:
    chr_val = expr["chr"]
    items = expr.get("items", [])
    if chr_val in {"^", "ˆ", "\u0302"}:
        return "\u0302"
    if chr_val in {"→", "⃗", "\u20d7"}:
        if len(items) == 1 and str(items[0]) in {"i", "j", "k"}:
            return "\u0302"
        return "\u20d1"
    return chr_val


def nary_fragment(chr_val: str, sub: list[Any], sup: list[Any], body: list[Any]) -> str:
    return (
        "<m:nary><m:naryPr>"
        f'<m:chr m:val="{escape(chr_val)}"/><m:limLoc m:val="undOvr"/>'
        "</m:naryPr><m:sub>" + items_fragment(sub) + "</m:sub><m:sup>"
        + items_fragment(sup) + "</m:sup><m:e>" + items_fragment(body)
        + "</m:e></m:nary>"
    )


def binom_fragment(top: list[Any], bottom: list[Any]) -> str:
    return delim(
        [{"kind": "matrix", "brackets": "none", "rows": [[top], [bottom]]}],
        "(",
        ")",
    )


def limit_fragment(expr: dict[str, Any]) -> str:
    var = expr.get("var", expr.get("base_var", expr.get("x", [])))
    to = expr.get("to", expr.get("target", []))
    if "lim" in expr:
        lim = expr["lim"]
    else:
        lim = []
        if var not in (None, "", []):
            lim.extend(var if isinstance(var, list) else [var])
        if to not in (None, "", []):
            if lim:
                lim.append("→")
            lim.extend(to if isinstance(to, list) else [to])
    body = expr.get("body", [])
    return (
        "<m:limLow><m:e>" + mtext("lim") + "</m:e><m:lim>" + items_fragment(lim)
        + "</m:lim></m:limLow>" + items_fragment(body)
    )


def item_to_omml_fragment(item: Any) -> str:
    if isinstance(item, dict):
        return expr_fragment(item)
    value = str(item)
    if contains_thai(value):
        raise ValueError(f"Thai text leaked into generic math item: {value!r}")
    if value in VARIABLES:
        return mr(value)
    if value in FUNCTION_NAMES:
        return mtext(value)
    if value in OPERATOR_TOKENS or value in TIGHT_PREFIX_OPERATORS:
        return mop(value)
    if value == ",":
        return mtext(", ")
    # Composite/implicit-product string such as "3x", "-2x", "ac", "2π", "x_1":
    # decompose with the one shared grammar so a variable adjacent to a
    # coefficient or another variable is still italic. Idempotent for atoms.
    return compact_item_to_omml_fragment(value)


def compact_item_to_omml_fragment(value: str) -> str:
    tokens = normalize_math_string(value)
    if tokens != [value]:
        # grammar split it (e.g. "3x" -> ["3","x"], "x_1" -> [sub-dict]); render each
        return "".join(item_to_omml_fragment(token) for token in tokens)
    # one atomic token that matched no whitelist above: a product of single-letter
    # variables like "ac"/"xy". Known functions are handled before we get here, so
    # a bare multi-letter alphabetic run is variables. Deliberate upright
    # multi-letter identifiers/units/labels must be passed as {"kind": "upright"}.
    if value.isalpha() and len(value) > 1:
        return "".join(mr(char) for char in value)
    return mtext(value)


def expr_fragment(expr: dict[str, Any]) -> str:
    kind = expr["kind"]
    if kind == "plain":
        return item_to_omml_fragment(expr["value"])
    if kind == "expr":
        return items_fragment(expr["items"])
    if kind == "thai_text":
        return thai_mtext(expr["text"])
    if kind == "upright":
        return mtext(expr["text"])
    if kind in {"paren", "delim"}:
        return delim(expr["items"], expr.get("beg", "("), expr.get("end", ")"))
    if kind == "neg":
        return mop("−") + items_fragment(expr["items"])
    if kind == "sup":
        return "<m:sSup><m:e>" + items_fragment(expr["base"]) + "</m:e><m:sup>" + items_fragment(expr["sup"]) + "</m:sup></m:sSup>"
    if kind == "sub":
        return "<m:sSub><m:e>" + items_fragment(expr["base"]) + "</m:e><m:sub>" + items_fragment(expr["sub"]) + "</m:sub></m:sSub>"
    if kind == "sub_sup":
        return (
            "<m:sSubSup><m:e>" + items_fragment(expr["base"]) + "</m:e><m:sub>"
            + items_fragment(expr["sub"]) + "</m:sub><m:sup>" + items_fragment(expr["sup"])
            + "</m:sup></m:sSubSup>"
        )
    if kind == "frac":
        return "<m:f><m:num>" + items_fragment(expr["num"]) + "</m:num><m:den>" + items_fragment(expr["den"]) + "</m:den></m:f>"
    if kind == "rad":
        if "deg" in expr:
            return "<m:rad><m:deg>" + items_fragment(expr["deg"]) + "</m:deg><m:e>" + items_fragment(expr["items"]) + "</m:e></m:rad>"
        return '<m:rad><m:radPr><m:degHide m:val="on"/></m:radPr><m:deg/><m:e>' + items_fragment(expr["items"]) + "</m:e></m:rad>"
    if kind == "bar":
        return '<m:bar><m:barPr><m:pos m:val="top"/></m:barPr><m:e>' + items_fragment(expr["items"]) + "</m:e></m:bar>"
    if kind == "acc":
        chr_val = accent_character(expr)
        return f'<m:acc><m:accPr><m:chr m:val="{escape(chr_val)}"/></m:accPr><m:e>' + items_fragment(expr["items"]) + "</m:e></m:acc>"
    if kind == "matrix":
        return bracketed_matrix_fragment(
            expr["rows"],
            expr.get("brackets", "[]"),
            expr.get("col_aligns"),
            expr.get("base_jc"),
        )
    if kind == "func":
        return "<m:func><m:fName>" + mtext(expr["name"]) + "</m:fName><m:e>" + items_fragment(expr["arg"]) + "</m:e></m:func>"
    if kind == "log":
        return (
            "<m:func><m:fName><m:sSub><m:e>" + mtext("log") + "</m:e><m:sub>"
            + items_fragment(expr["base"]) + "</m:sub></m:sSub></m:fName><m:e>"
            + items_fragment(expr["arg"]) + "</m:e></m:func>"
        )
    if kind == "lim":
        return limit_fragment(expr)
    if kind == "lim_low":
        return "<m:limLow><m:e>" + items_fragment(expr["base"]) + "</m:e><m:lim>" + items_fragment(expr["lim"]) + "</m:lim></m:limLow>"
    if kind == "nary":
        return nary_fragment(expr["chr"], expr["sub"], expr["sup"], expr["body"])
    if kind == "integral":
        return nary_fragment(
            expr.get("chr", "∫"),
            expr.get("sub", expr.get("from", [])),
            expr.get("sup", expr.get("to", [])),
            expr.get("body", expr.get("items", [])),
        )
    if kind == "binom":
        return binom_fragment(
            expr.get("top", expr.get("n", expr.get("upper", []))),
            expr.get("bottom", expr.get("k", expr.get("lower", []))),
        )
    if kind == "cases":
        col_aligns = expr.get("col_aligns") or ["left", "left"]
        base_jc = expr.get("base_jc") or "left"
        return (
            '<m:d><m:dPr><m:begChr m:val="{"/><m:endChr m:val=""/></m:dPr><m:e>'
            + matrix_fragment(expr["rows"], col_aligns, base_jc)
            + "</m:e></m:d>"
        )
    raise ValueError(f"Unsupported OMML expression kind: {kind}")


def math_omml(expr: dict[str, Any] | list[Any] | str) -> str:
    if isinstance(expr, dict):
        body = expr_fragment(expr)
    elif isinstance(expr, list):
        body = items_fragment(expr)
    else:
        body = item_to_omml_fragment(expr)
    return f'<m:oMath xmlns:m="{M_NS}" xmlns:w="{W_NS}">{body}</m:oMath>'


def append_math(paragraph: Any, expr: dict[str, Any] | list[Any] | str) -> None:
    paragraph._p.append(parse_xml(math_omml(expr)))


def ensure_thai_insertion_safe_paragraph_end(paragraph: Any) -> None:
    """Close a paragraph that ends with an equation with a Thai body run.

    Word inherits formatting from the run to the left of the cursor. A
    paragraph whose last child is ``m:oMath`` leaves an OMML run there, so an
    empty insertion-safe Thai run is appended for manual typing to inherit.
    """
    trailing_tags = (qn("w:r"), qn("m:oMath"), qn("m:oMathPara"))
    trailing = [child for child in paragraph._p if child.tag in trailing_tags]
    if not trailing or trailing[-1].tag == qn("w:r"):
        return
    set_thai_body_run(paragraph.add_run())


def append_parts(paragraph: Any, parts: list[dict[str, Any]]) -> None:
    for part in parts:
        part_type = part["type"]
        if part_type == "text":
            run = paragraph.add_run(part["text"])
            set_thai_body_run(run, bold=part.get("bold"))
        elif part_type == "latin_text":
            run = paragraph.add_run(part["text"])
            set_latin_run(run, bold=part.get("bold"))
        elif part_type == "label":
            run = paragraph.add_run(part["text"])
            set_thai_label_run(run, bold=part.get("bold", True), size=part.get("size", 16))
        elif part_type == "math":
            append_math(paragraph, part.get("expr", part))
        elif part_type == "line_break":
            paragraph.add_run().add_break()
        else:
            raise ValueError(f"Unsupported part type: {part_type}")
    ensure_thai_insertion_safe_paragraph_end(paragraph)


def append_parts_or_tables(doc: Document, paragraph: Any, parts: list[dict[str, Any]], space_after: int = 4) -> Any:
    """Append inline parts while routing table blocks through ``add_table``.

    ``python-docx`` paragraphs cannot contain real Word tables. Shared year
    builders previously duplicated this routing locally; keeping it here makes
    table-capable sources portable across years.
    """
    inline: list[dict[str, Any]] = []
    for part in parts:
        if part.get("type") == "table":
            if inline:
                append_parts(paragraph, inline)
                inline = []
            add_table(
                doc,
                part["rows"],
                part.get("widths"),
                layout_profile=part.get("layout_profile"),
            )
            paragraph = doc.add_paragraph()
            configure_paragraph(paragraph, space_after=space_after)
        else:
            inline.append(part)
    if inline:
        append_parts(paragraph, inline)
    return paragraph


def add_paragraph(doc: Document, parts: list[dict[str, Any]] | None = None, space_after: int = 4) -> Any:
    paragraph = doc.add_paragraph()
    configure_paragraph(paragraph, space_after=space_after)
    if parts:
        append_parts_or_tables(doc, paragraph, parts, space_after=space_after)
    return paragraph


def add_heading(doc: Document, text: str, space_after: int = 6) -> Any:
    return add_paragraph(doc, [{"type": "label", "text": text, "bold": True}], space_after=space_after)


def add_question_block(doc: Document, number: int, prompt_parts: list[dict[str, Any]], space_after: int = 6) -> None:
    paragraph = doc.add_paragraph()
    configure_paragraph(paragraph, space_after=space_after)
    label = paragraph.add_run(f"ข้อ {number}. ")
    set_thai_label_run(label, bold=True)
    append_parts_or_tables(doc, paragraph, prompt_parts, space_after=space_after)


def set_table_fixed_widths(table: Any, widths: list[float]) -> None:
    """Compatibility wrapper for callers whose explicit widths are in inches."""
    set_table_fixed_widths_cm(table, [width * 2.54 for width in widths])


def standard_activity_table_widths(column_count: int) -> list[float]:
    """Return the current default, which is only the 16 cm one-column layout."""
    if column_count != 1:
        raise ValueError(
            "multi-column student tables require an explicit layout profile or explicit widths"
        )
    return [STANDARD_ACTIVITY_TABLE_WIDTH_IN]


def current_student_table_widths(layout: str) -> list[float]:
    """Return current student-table profile widths in inches for ``add_table``."""
    profile = get_current_layout_profile("student-question-layout")
    try:
        widths_cm = profile["table_widths_cm"][layout]
    except KeyError as exc:
        raise ValueError(f"unsupported student table layout: {layout!r}") from exc
    return [width / 2.54 for width in widths_cm]


def add_table(
    doc: Document,
    rows: list[list[list[dict[str, Any]]]],
    widths: list[float] | None = None,
    *,
    layout_profile: str | None = None,
) -> Any:
    if widths is not None and layout_profile is not None:
        raise ValueError("pass either widths or layout_profile, not both")
    if layout_profile is not None:
        widths = current_student_table_widths(layout_profile)
    elif widths is None:
        widths = standard_activity_table_widths(len(rows[0]))
    if len(widths) != len(rows[0]):
        raise ValueError("width count must match table column count")
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_fixed_widths(table, widths)
    for row_index, row in enumerate(rows):
        for col_index, cell_parts in enumerate(row):
            cell = table.cell(row_index, col_index)
            cell.width = Inches(widths[col_index])
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            configure_paragraph(para, space_after=0)
            append_parts(para, cell_parts)
    return table


def save_docx(doc: Document, path: str | Path) -> Path:
    out = Path(path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    normalize_docx_theme_thai_fonts(out)
    return out


def normalize_docx_theme_thai_fonts(path: str | Path, target_font: str = "TH Sarabun New") -> None:
    """Set Thai theme mappings in saved DOCX packages.

    ``python-docx`` starts from Word's default theme, which maps Thai to
    Angsana/Cordia. The post-build thai-font-normalize gate repairs this, but
    generated documents should be correct before the repair layer runs.
    """
    path = Path(path)
    ET.register_namespace("a", A_NS)
    changed = False
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp_path = Path(tmp.name)
    try:
        with zipfile.ZipFile(path, "r") as zin, zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
            for info in zin.infolist():
                data = zin.read(info.filename)
                if info.filename.startswith("word/theme/") and info.filename.endswith(".xml"):
                    root = ET.fromstring(data)
                    file_changed = False
                    for font in root.findall(".//{" + A_NS + "}font"):
                        if font.get("script") == "Thai" and font.get("typeface") != target_font:
                            font.set("typeface", target_font)
                            file_changed = True
                    if file_changed:
                        data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
                        changed = True
                zout.writestr(info, data)
        if changed:
            tmp_path.replace(path)
        else:
            tmp_path.unlink(missing_ok=True)
    except Exception:
        tmp_path.unlink(missing_ok=True)
        raise

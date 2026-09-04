"""S07: repairing an imported / legacy-font Thai .docx.

thai-docx repairs an imported document by shelling out to the shared
`thai-font-normalize/scripts/fix-thai-font`, which routes Thai Complex Script to
TH Sarabun New (and fixes the theme). This is the sibling-by-absolute-path reuse
pattern; thai-docx owns no font-remap code of its own.

Note: fix-thai-font deliberately fixes only the Complex-Script (`w:cs`) slot — that
is what makes Thai render correctly — and leaves Latin slots as authored (so
thai-math-docx keeps Latin=Cambria). A legacy font left in a Latin slot is a
cosmetic follow-up, not a Thai-rendering defect.
"""
from __future__ import annotations

from pathlib import Path
import collections
import re
import subprocess
import unittest
import zipfile

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SKILLS_ROOT = Path(__file__).resolve().parents[2]
sys_scripts = SKILLS_ROOT / "thai-docx" / "scripts"
import sys
if str(sys_scripts) not in sys.path:
    sys.path.insert(0, str(sys_scripts))
import repair as repair_mod


def _all_slot_fonts(docx: Path) -> collections.Counter:
    xml = zipfile.ZipFile(docx).read("word/document.xml").decode("utf-8", "replace")
    return collections.Counter(re.findall(r'w:(?:ascii|hAnsi|cs|eastAsia)="([^"]+)"', xml))


def _make_imported_doc(path: Path) -> None:
    doc = Document()
    # Thai run declaring a legacy font in every slot
    run = doc.add_paragraph().add_run("เอกสารนำเข้าที่ประกาศฟอนต์ TH SarabunPSK ทุก slot")
    rpr = run._r.get_or_add_rPr()
    rfonts = OxmlElement("w:rFonts")
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), "TH SarabunPSK")
    rpr.append(rfonts)
    # a numbers-only run on the legacy font in the Latin slot (the budget-doc case)
    num = doc.add_paragraph().add_run("13,718,220")
    nrp = num._r.get_or_add_rPr()
    nrf = OxmlElement("w:rFonts")
    for attr in ("w:ascii", "w:hAnsi"):
        nrf.set(qn(attr), "TH SarabunPSK")
    nrp.append(nrf)
    # a genuine Latin font that must be PRESERVED
    lat = doc.add_paragraph().add_run("Calibri stays")
    lrp = lat._r.get_or_add_rPr()
    lrf = OxmlElement("w:rFonts")
    for attr in ("w:ascii", "w:hAnsi"):
        lrf.set(qn(attr), "Calibri")
    lrp.append(lrf)
    doc.save(path)


class RepairImportedTests(unittest.TestCase):
    def test_repair_converts_every_legacy_slot_to_new_and_keeps_real_latin(self) -> None:
        import tempfile
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "imported-psk.docx"
            _make_imported_doc(path)
            before = _all_slot_fonts(path)
            self.assertIn("TH SarabunPSK", before)

            result = repair_mod.repair(path, backup=False)

            after = _all_slot_fonts(path)
            # no legacy Thai font left in ANY slot (Thai cs + numbers ascii both fixed)
            self.assertNotIn("TH SarabunPSK", after, msg=f"legacy font remained: {after}")
            self.assertIn("TH Sarabun New", after)
            # a genuine Latin font is left untouched
            self.assertIn("Calibri", after, msg="repair must not clobber real Latin fonts")
            self.assertGreater(result["legacy_slots_swept"], 0)


if __name__ == "__main__":
    unittest.main()
